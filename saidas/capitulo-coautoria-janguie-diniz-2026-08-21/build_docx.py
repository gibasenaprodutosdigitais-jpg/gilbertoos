#!/usr/bin/env python3
"""
Gera o capítulo em .docx formatado nas normas ABNT:
- Times New Roman 12
- Espaçamento 1,5
- Margens: 3cm superior/esquerda, 2cm inferior/direita
- Justificado, recuo de parágrafo 1,25cm
- Numeração de página no canto superior direito
"""
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = Pt(12)

doc = Document()

# ---------- Margens ----------
section = doc.sections[0]
section.top_margin = Cm(3)
section.left_margin = Cm(3)
section.bottom_margin = Cm(2)
section.right_margin = Cm(2)

# ---------- Estilo padrão ----------
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = SIZE
# garante fonte correta também para caracteres de outros scripts (acentuação)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), FONT)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = FONT
    run.font.size = SIZE


# ---------- Cabeçalho: numeração no canto superior direito ----------
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
add_page_number_field(hp)


def body(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = SIZE
    return p


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.bold = True
    return p


def author(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = SIZE
    run.italic = True
    return p


def epigraph(lines, source):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(lines)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.left_indent = Cm(8)
    p2.paragraph_format.space_after = Pt(24)
    run2 = p2.add_run(source)
    run2.font.name = FONT
    run2.font.size = Pt(10)
    run2.italic = True
    return p2


def heading(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{number} {text}" if number else text)
    run.font.name = FONT
    run.font.size = SIZE
    run.font.bold = True
    return p


# =========================================================
# CONTEÚDO
# =========================================================

title("Gestão: o Alicerce Invisível da Nova Economia")
author("Gilberto Sena")

epigraph(
    '"Os 5 Gs da gestão, que detalhei ao longo dessas páginas não são meras '
    'teorias, até porque, eu não sou um acadêmico, não sou um homem de '
    'teorias. Pelo contrário, são, de fato, a essência da minha prática."',
    "— Os 5 G's da Gestão Empreendedora",
)

heading("", "Introdução")

body(
    "Há pouco mais de doze anos, eu estava numa moto, entregando encomenda, "
    "correndo contra o relógio pra fechar a meta do dia. Antes disso, fui "
    "servente de pedreiro, fui camelô, tive um carrinho de cachorro-quente "
    "na rua. Hoje sou fundador e CEO do Grupo Sena Soluções Empresariais, "
    "com sede em Belo Horizonte e atuação em seis estados brasileiros, e "
    "assino, com orgulho, um capítulo ao lado de um dos maiores nomes "
    "empresariais do país sobre o tema que mais me apaixona: como "
    "construir negócios que duram, escalam e transformam."
)
body(
    "Conto essa trajetória não por vaidade, mas porque ela é o ponto de "
    "partida obrigatório deste texto. Quando me convidaram para escrever "
    "sobre a Nova Economia, a expectativa implícita, eu sei, era de que um "
    "profissional da área contábil e tributária escrevesse sobre "
    "conformidade, sobre obrigação acessória, sobre a letra fria da lei. "
    "Eu escolhi um caminho diferente. Escolho reinterpretar a minha "
    "profissão — e não abandoná-la — dentro dos cinco pilares que "
    "estruturam esta obra: influência, inovação, criação de valor, "
    "escalabilidade e impacto."
)
body(
    "Minha tese é simples de enunciar e difícil de aceitar por quem ainda "
    "enxerga a gestão financeira e tributária como departamento de "
    "suporte: nenhum negócio da Nova Economia — por mais disruptivo, "
    "digital ou influente que seja — sobrevive, escala ou gera impacto "
    "real sem uma estrutura de gestão sólida por trás. A tecnologia muda "
    "o produto. A gestão sustenta a empresa. E é exatamente aí, no "
    "alicerce invisível, que proponho que o leitor me acompanhe."
)
body(
    "No livro que escrevi, Os 5 G's da Gestão Empreendedora, sistematizei "
    "cinco frentes que considero inegociáveis para qualquer empreendedor: "
    "Gestão de Si, Gestão Financeira, Gestão de Pessoas, Gestão "
    "Tributária e Gestão de Fornecedores e Clientes. Não são teoria de "
    "sala de aula — são a essência de trinta anos de prática, incluindo "
    "sete falências que enfrentei ao longo do caminho e das quais me "
    "reergui. É esse arcabouço que vou usar, ao longo das próximas "
    "páginas, para atravessar os cinco pilares desta obra."
)

heading("1.", "Influência: de quem sabe a lei para quem constrói autoridade")

body(
    "A Nova Economia redefiniu o que significa ser influente. Não se "
    "trata mais de quem tem o maior cargo ou o diploma mais respeitado — "
    "trata-se de quem consegue, de fato, mudar a decisão de outra "
    "pessoa. Sob essa definição, influência deixou de ser um ativo de "
    "artistas e criadores de conteúdo para se tornar um ativo de "
    "qualquer profissional que consiga traduzir complexidade em direção "
    "clara."
)
body(
    "Na minha área, essa mudança é radical. Durante décadas, o contador "
    "e o consultor tributário foram vistos como figuras de bastidor — "
    "quem resolve o problema depois que ele já aconteceu, quem cuida da "
    "papelada enquanto o empresário \"toca o negócio de verdade\". Eu me "
    "recuso a operar dentro dessa moldura. A diferença entre o contador "
    "tradicional e o gestor que eu me proponho a ser é exatamente essa: "
    "o contador diz o que a lei manda; o gestor diz o que fazer com "
    "isso. Um entrega informação. O outro entrega decisão."
)
body(
    "Essa mudança de postura é, ela própria, um exercício de influência. "
    "Quando um empresário chega até mim assustado com uma nova exigência "
    "fiscal e sai da conversa entendendo não apenas a obrigação, mas o "
    "que aquilo significa para o preço do seu produto, para o fluxo de "
    "caixa do próximo trimestre, para a decisão de contratar ou não — eu "
    "não apenas informei: eu influenciei uma decisão estratégica. "
    "Multiplicado por milhares de empresários ao longo de mais de uma "
    "década à frente do Grupo Sena, esse tipo de influência silenciosa "
    "tem um peso econômico que raramente é contabilizado nas métricas de "
    "alcance da Nova Economia, mas que sustenta boa parte dela."
)
body(
    "Há uma frase que resume minha filosofia sobre o assunto tributário, "
    "e que uso com frequência: no veneno está o antídoto. Não aconselho, "
    "em hipótese alguma, que um empresário tente contornar ou burlar a "
    "legislação — o que parece economia hoje pode se tornar prejuízo "
    "grande amanhã. Mas o mesmo conhecimento técnico que assusta o "
    "empresário desorganizado é, nas mãos certas, a ferramenta que o "
    "protege e o faz crescer. Transformar veneno em antídoto, através de "
    "conhecimento aplicado, é a forma mais honesta de influência que "
    "conheço — porque não depende de audiência, depende de resultado."
)
body(
    "Também construí, ao longo dos anos, um canal de conteúdo digital "
    "que começou em 2013 e hoje alcança centenas de milhares de pessoas, "
    "onde traduzo, quase diariamente, o que está mudando na legislação "
    "tributária brasileira para uma linguagem que qualquer empresário "
    "entende. Não faço isso para acumular seguidores. Faço isso porque "
    "entendi, cedo, que na Nova Economia a autoridade técnica sem "
    "tradução acessível é autoridade desperdiçada. Influência real não é "
    "o tamanho do público — é a distância que se consegue percorrer "
    "entre o conhecimento complexo e a decisão prática de quem precisa "
    "dele."
)

heading("2.", "Inovação: o método como forma de inovar num setor conservador")

body(
    "Quando se fala em inovação na Nova Economia, o imaginário corre "
    "imediatamente para startups, inteligência artificial, novos "
    "modelos de plataforma. Raramente alguém associa inovação a "
    "contabilidade, tributação ou gestão financeira. Essa é, a meu ver, "
    "uma leitura incompleta do que significa inovar."
)
body(
    "Inovação não é apenas criar um produto novo. É, também, criar um "
    "método novo para resolver um problema antigo. E poucos setores no "
    "Brasil carregam um problema tão antigo, tão mal resolvido e tão "
    "urgente quanto a gestão tributária das empresas brasileiras — um "
    "problema que, neste exato momento, está sendo redesenhado pela "
    "reforma tributária em curso, com a substituição de cinco tributos "
    "por um modelo de IVA dual, a extinção de PIS e Cofins, a chegada do "
    "split payment e uma transição que se estende até 2033. Trata-se, "
    "sem exagero, do maior evento de inovação forçada que o mercado "
    "empresarial brasileiro já viveu — e a maioria dos empresários ainda "
    "não entendeu o tamanho da mudança que já começou."
)
body(
    "Minha contribuição pessoal a esse cenário não foi inventar uma nova "
    "tecnologia. Foi sistematizar um método. Ao longo da minha "
    "trajetória, desenvolvi frameworks próprios que hoje aplico com "
    "clientes e ensino a outros gestores: a Escada do Engajamento, que "
    "organiza em seis degraus o caminho entre o conhecimento de um "
    "colaborador sobre a empresa e seu comprometimento real com ela; a "
    "Matriz da Gestão da Comunicação, que obriga qualquer comunicado "
    "interno a responder seis perguntas antes de ser enviado; o TRI — "
    "Triagem, Resposta, Inovação —, framework que desenvolvi para gerir "
    "risco de fornecedores e clientes; e o Geo Ciclo do Grupo Sena, "
    "metodologia própria para estruturar e reestruturar negócios de "
    "qualquer porte."
)
body(
    "Nenhum desses frameworks nasceu em uma sala de aula. Nasceram da "
    "necessidade prática de resolver, repetidamente, os mesmos problemas "
    "que via em empresas diferentes — e da constatação de que ensinar "
    "por metodologia, não por improviso, é o que permite que o "
    "conhecimento se multiplique além da minha própria agenda. Essa é, "
    "para mim, a forma mais honesta de inovação num setor "
    "tradicionalmente conservador: não reinventar a lei, mas reinventar "
    "a forma como o empresário se relaciona com ela."
)
body(
    "Há também uma dimensão de inovação na própria filosofia de "
    "negócio. Substituí, na gestão de pessoas do Grupo Sena, a avaliação "
    "de desempenho formal e burocrática por uma metodologia que chamo de "
    "Feedback Alavancado — uma conversa contínua, de abordagem "
    "narrativa, que conta a história do que aconteceu em vez de "
    "simplesmente apontar o defeito. É uma inovação pequena na "
    "superfície, mas profunda no resultado: transforma feedback em "
    "construção, não em veredito. E é exatamente esse tipo de inovação "
    "— silenciosa, metodológica, replicável — que sustenta negócios "
    "capazes de crescer sem perder a alma pelo caminho."
)

heading("3.", "Criação de valor: o dinheiro como professor, não como fim")

body(
    "Um dos enganos mais recorrentes da Nova Economia é tratar criação "
    "de valor como sinônimo de valuation, de captação, de múltiplo sobre "
    "receita. Esses números importam, mas são consequência — não são a "
    "origem do valor. A origem está em algo mais simples e mais difícil: "
    "saber, de fato, para onde vai o dinheiro de um negócio, e o que "
    "fazer com essa informação."
)
body(
    "Escrevi, no meu livro, que o dinheiro nunca foi para mim apenas uma "
    "ferramenta, mas um professor. Ele me ensinou disciplina quando "
    "precisei economizar, inovação quando precisei investir sem "
    "certeza, e humildade quando enfrentei desafios financeiros — e "
    "posso falar disso com propriedade, porque ao longo da minha "
    "trajetória enfrentei a falência não uma, mas sete vezes. Cada uma "
    "delas foi uma lição dura, mas também uma oportunidade de "
    "aprendizado. Não escondo essas quedas porque elas são, precisamente, "
    "a fonte da minha autoridade para falar sobre criação de valor: "
    "quem nunca perdeu dinheiro dificilmente entende, na prática, o que "
    "significa protegê-lo."
)
body(
    "Na estrutura que desenvolvi para a Gestão Financeira, organizo o "
    "tema em formato de pirâmide: Governança Financeira na base, seguida "
    "por Gestão Orçamentária, Gestão de Fluxo de Caixa, Gestão de "
    "Investimentos e, no topo, Gestão de Endividamento. A ordem não é "
    "arbitrária. A maioria dos empresários que atendo quer pular direto "
    "para investimento sem ter governança na base — é como construir "
    "cobertura sem fundação. Valor real não se cria pulando etapas; "
    "cria-se subindo a pirâmide na ordem certa."
)
body(
    "É por essa razão que considero a gestão tributária eficiente uma "
    "das formas mais subestimadas de criação de valor na economia "
    "brasileira. Cada real de imposto pago corretamente, sem desperdício "
    "por desconhecimento, sem multa por atraso, sem enquadramento "
    "tributário inadequado, é um real que permanece disponível para "
    "reinvestimento, contratação, inovação. Ao longo de mais de uma "
    "década à frente do Grupo Sena, ajudamos empresas de diferentes "
    "portes a recuperar e a preservar centenas de milhões de reais que, "
    "de outra forma, teriam se perdido em ineficiência tributária — não "
    "por sonegação, mas por planejamento. Essa é uma forma de criação de "
    "valor que não aparece em pitch deck, mas que decide se uma empresa "
    "sobrevive ao próximo ciclo econômico."
)

heading("4.", "Escalabilidade: sistemas que não dependem do fundador")

body(
    "Todo empreendedor, em algum momento, confunde crescimento com "
    "escala. Eu mesmo cometi esse erro. Crescer é vender mais. Escalar "
    "é vender mais sem que a estrutura da empresa cresça na mesma "
    "proporção — o que só é possível quando existem sistemas, e não "
    "apenas pessoas talentosas, segurando o negócio de pé."
)
body(
    "Um dos textos mais provocadores que escrevi no meu livro pergunta "
    "ao leitor: se eu te perguntasse qual é o propósito da sua empresa, "
    "o que você diria? A resposta que mais ouço é: crescer. E, se essa "
    "também for a resposta do leitor, eu tenho uma notícia difícil: sua "
    "empresa não é diferente de um tumor. Crescimento sem propósito é "
    "crescimento descontrolado. Crescimento não é propósito — é "
    "consequência do propósito, e escala não é sinônimo de tamanho — é "
    "sinônimo de sistema."
)
body(
    "A escalabilidade, na minha experiência prática, se constrói em "
    "três camadas. A primeira é a matriz de responsabilidades: garantir "
    "que nenhuma tarefa crítica da empresa dependa de uma única pessoa, "
    "incluindo o próprio fundador. A segunda é a instrução de trabalho: "
    "transformar o conhecimento tácito de quem \"sabe fazer\" em processo "
    "documentado que qualquer pessoa treinada consegue executar com o "
    "mesmo padrão. A terceira, a mais difícil emocionalmente, é a "
    "delegação deliberada: separar, todos os dias, o que é tarefa "
    "operacional do que é decisão estratégica, e devolver ao gestor "
    "apenas a segunda."
)
body(
    "O próprio Grupo Sena é prova viva dessa lógica. Comecei sozinho, em "
    "2010, saindo de uma operadora de telefonia onde já resolvia "
    "problemas de empresas grandes como funcionário — decidi que "
    "resolveria os mesmos problemas através do meu próprio negócio. "
    "Hoje, mais de uma década depois, o grupo opera com estrutura em "
    "seis estados brasileiros, sem que cada decisão precise passar pela "
    "minha mesa. Isso não aconteceu por sorte. Aconteceu porque apliquei, "
    "na minha própria empresa, o mesmo método que hoje ensino: o Geo "
    "Ciclo, minha metodologia própria para estruturar e reestruturar "
    "negócios, nasceu exatamente da necessidade de sistematizar o que "
    "antes vivia só na minha cabeça."
)
body(
    "Na Nova Economia, onde modelos de negócio digitais prometem escala "
    "instantânea, vale lembrar que toda escala sustentável — digital ou "
    "não — depende de uma retaguarda financeira, tributária e de gestão "
    "de pessoas que suporte o volume sem quebrar. Já vi negócios "
    "crescerem rápido demais e afundarem justamente por não terem essa "
    "retaguarda. Escala sem gestão não é crescimento — é contagem "
    "regressiva."
)

heading("5.", "Impacto: transformar dor vivida em legado")

body(
    "Se influência, inovação, criação de valor e escalabilidade são os "
    "meios, impacto é o motivo que sustenta o esforço no longo prazo. E "
    "aqui volto, deliberadamente, à moto onde comecei este texto."
)
body(
    "Fui forjado pelos momentos que passei. Fui servente de pedreiro, "
    "fui camelô, tive carrinho de cachorro-quente, fui motoboy — e cada "
    "uma dessas experiências construiu a base de negócio e de rede de "
    "relações que sustenta o que faço hoje. Não conto essa história "
    "para inspirar por inspirar. Conto porque ela define o tipo de "
    "impacto que persigo: eu me orgulho de saber que existem muitos "
    "\"Gilbertos\" espalhados por aí, gente que só precisa de "
    "conhecimento e apoio para crescer."
)
body(
    "Impacto, na Nova Economia, é frequentemente medido por métricas de "
    "alcance — visualizações, seguidores, engajamento. São métricas "
    "válidas, mas superficiais quando isoladas. O impacto que realmente "
    "me interessa é outro: quantas empresas continuam de pé porque "
    "entenderam a tempo uma mudança tributária; quantos empregos foram "
    "preservados porque um empresário aceitou fazer a gestão financeira "
    "antes de ser tarde demais; quantas famílias mantiveram sua "
    "estabilidade porque alguém, em algum momento, transformou um "
    "problema técnico complexo em uma decisão simples e a tempo."
)
body(
    "A missão que orienta o Grupo Sena, e que assino com a mesma "
    "convicção de quando a escrevemos pela primeira vez, resume esse "
    "compromisso: atuar no mercado corporativo de assessoria e "
    "consultoria contábil, tributária, jurídica e de gestão financeira "
    "com o principal objetivo de garantir a tranquilidade de nossos "
    "clientes, de forma ética, segura e confiável. Não é uma missão "
    "feita para caber em um slide de apresentação. É uma missão feita "
    "para orientar decisões nos dias em que ninguém está olhando."
)
body(
    "Quando eu vendia cachorro-quente, quando eu era motoboy, eu não "
    "tinha ideia de onde a vida me levaria — mas tinha o coração cheio "
    "de sonhos e a determinação de transformar a minha realidade. Se "
    "este capítulo tiver algum valor para o leitor que atua em qualquer "
    "área da Nova Economia, que seja este: acredite em si mesmo quando "
    "as adversidades tentarem te derrubar. Eu acreditei, mesmo quando "
    "fui questionado e testado pelas sete quedas que enfrentei. O mundo "
    "dos negócios aguarda as conquistas de quem lê este capítulo. Que "
    "transformemos, juntos, sonhos em realidades tangíveis, e deixemos "
    "um legado que transcenda gerações."
)

heading("", "Considerações finais")

body(
    "A Nova Economia costuma ser narrada pela lente de quem cria "
    "produtos visíveis: aplicativos, plataformas, marcas pessoais, "
    "conteúdo. Este capítulo defendeu uma tese diferente, construída a "
    "partir da minha própria trajetória e do método que sistematizei ao "
    "longo de mais de uma década à frente do Grupo Sena: por trás de "
    "todo negócio influente, inovador, gerador de valor, escalável e "
    "impactante, existe uma estrutura de gestão — financeira, "
    "tributária, de pessoas — que raramente aparece nos holofotes, mas "
    "sem a qual nenhum dos cinco pilares se sustenta."
)
body(
    "Reinterpretar a própria profissão dentro da Nova Economia não "
    "significa abandonar sua essência técnica. Significa compreender "
    "que técnica sem tradução não influencia, técnica sem método não "
    "inova, técnica sem aplicação não cria valor, técnica sem sistema "
    "não escala, e técnica sem propósito não gera impacto algum. Foi "
    "essa a jornada que tentei documentar aqui: da moto até a mesa onde "
    "hoje ajudo empresários a tomar as decisões mais importantes dos "
    "seus negócios. Que este capítulo sirva de convite a outros "
    "profissionais de formação dita conservadora para que também "
    "reconheçam, na própria prática cotidiana, os cinco pilares da "
    "economia que estamos, juntos, construindo."
)

heading("", "Sobre o autor")

body(
    "Gilberto Luís de Sena é fundador e CEO do Grupo Sena Soluções "
    "Empresariais, referência em contabilidade, auditoria fiscal e "
    "tributária, assessoria jurídica e gestão financeira, com sede "
    "principal em Belo Horizonte e atuação em Goiânia, Parauapebas, São "
    "Paulo, Brasília e Vitória. É Bacharel em Teologia pelo Seminário de "
    "Educação Teológica Kerigma Didache (SETEAD), pós-graduado em "
    "Direito Tributário (720 horas) pela FAVENI/Instituto de Educação "
    "Século XXI, e Agente Autônomo de Investimento (AAI) habilitado pela "
    "ANCORD. É autor do livro Os 5 G's da Gestão Empreendedora e criador "
    "do Geo Ciclo do Grupo Sena, metodologia própria de estruturação e "
    "reestruturação de negócios.",
    indent=False,
)

doc.save("capitulo-gestao-nova-economia.docx")
print("DOCX salvo.")
